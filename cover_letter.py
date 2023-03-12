from docx import Document

path = "C:\\Users\\Tanner\\Documents\\"
inputs = []

doc = Document("C:\\Users\\Tanner\\Documents\\Cover Letter.docx")

keywords = ["[Company]","[Position]"]

for keyword in keywords:
    input_i = input(keyword)
    inputs.append(input_i)

for i in range(len(keywords)):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if keywords[i] in paragraph.text:
                        paragraph.text = paragraph.text.replace(keywords[i], inputs[i])

docName = inputs[0] + " Cover Letter"

doc.save(path + docName + ".docx")
